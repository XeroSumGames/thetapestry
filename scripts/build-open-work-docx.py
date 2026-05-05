#!/usr/bin/env python
# build-open-work-docx.py
# Generate a print-friendly Word document of the open-work checklist.
#
# Why a script and not just a manual docx: the source-of-truth lives in
# tasks/open-work-2026-05-05-printable.md. When that markdown changes,
# re-run this to regenerate the .docx. Idempotent — overwrites the
# existing output file.
#
# Output: tasks/open-work-2026-05-05.docx
#
# Usage:
#   python scripts/build-open-work-docx.py
#
# Dependencies: python-docx (already installed on Tony's box).
#
# Design choices:
# - ☐ (Unicode 0x2610) prefix instead of Word content controls — content
#   controls print inconsistently across Word versions; the unicode box
#   is universal and crosses off cleanly with a pen.
# - 11pt body text + 0.75" margins → ~5 pages. Tighter than default.
# - Italic for design-call notes so they read as side-bar context.
# - Bold for "shipped" callouts on partial items (e.g. "(Recruit shipped)").

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# ── Content (source: tasks/open-work-2026-05-05-printable.md) ─────────
# Tuple format: (heading_level, text, optional_italic_note)
# heading_level: 0 = body, 1 = section H1, 2 = section H2, -1 = checkbox
# Special marker "PAGEBREAK" inserts a hard page break before the next
# section.

CONTENT = [
    # Title block — handled separately below

    (1, "TOP PRIORITY — From last night's playtest", None),
    (-1, "Perception check: redundant first modal. Should go straight to roll modal (auto-pick active PC).", None),
    (-1, "PCs riding Minnie don't move with her. Passengers in vehicle footprint should follow with offset preserved.",
        "Needs design call: stickiness vs. explicit mount/disembark; what happens on incompatible terrain."),
    (-1, "Random character — Medic produces no First Aid skill.",
        "Likely a wording mix-up — XSE has no \"First Aid\" skill; Medic seeds Medicine*. Confirm with player before chasing."),

    (1, "BUGS — Need repro / decision before code", None),
    (-1, "Initiative lag. Needs solo repro on your machine first.", None),
    (-1, "Damage calc spot-check. Reported \"2+2d6 (6) = 8 raw\" → should be 7/7. Replay verification.", None),
    (-1, "Failed skill checks still have two actions available. Code looks right; needs repro.", None),
    (-1, "Tactical map mouse-pan via drag — broken. WASD works; click-and-drag on empty cell doesn't. Multiple ship+revert attempts; \"no fix path identified\".", None),
    (-1, "HP render lag — previous-session follow-up.", None),

    (1, "PARTIALS — Finish what's started", None),
    (-1, "Modal unification. Normalize Stabilize, Distract, Coordinate, Group Check, Gut Instinct, First Impression to <RollModal>. (Recruit + Stress/Breaking/Lasting/Wound shipped.)", None),
    (-1, "Hide-NPCs reveal UX. Folder-level \"Reveal all in this folder\" + panic button \"reveal entire roster\". (Multi-select bar + auto-reveal-on-Start-Combat shipped.)", None),
    (-1, "Featured items. Thriver promote-to-featured for forum threads + war stories. (Module featured shipped.)", None),
    (-1, "DZ canon layer. District Zero-specific canon scope/UX. (Generic is_canon badge shipped.)", None),
    (-1, "DZ timeline visualization. Chronological page surfacing world-event timeline pins. (Timeline category + sort_order shipped.)", None),
    (-1, "Play stats per module. Track actuals — session count + avg player count. (Subscriber count shipped.)", None),
    (-1, "Sequence guards on loadRolls / loadChat. (loadEntries already has one.)", None),
    (-1, "Tier C1. Snapshot RPC for table-page mount. (Parallelization shipped via 96a66b2.)", None),
    (-1, "In-app SRD search. SRD copy is structurally complete; /rules/* just needs a search UI.", None),
    (-1, "Thriver godmode UI sweep. Widen `isGM &&` → `(isGM || isThriver)` in CampaignCommunity, CampaignObjects, VehicleCard, NpcRoster, character-sheet edits for non-owned PCs. (DB layer shipped.)", None),

    (1, "OLDER BUGS — Genuinely open", None),
    (-1, "Gut Instinct results presentation needs rework.",
        "Design discussion: narrative card vs. sheet overlay vs. GM DM."),
    (-1, "Inventory migration — auto-convert old string equipment to structured items on load.", None),
    (-1, "Allow characters in multiple campaigns.", None),
    (-1, "Transfer GM role; Session scheduling.", None),
    (-1, "Player-facing NPC card on Show All click. Currently opens GM-editable view; players should get read-only.", None),

    (1, "UX / POLISH", None),
    (-1, "Streamline logging into missions. /login → /stories → click → Join Session → /table is too many steps. Possible: deep-link, auto-redirect, \"Resume last session\" tile.", None),
    (-1, "King's Crossing Mall — tactical scenes (mall complex maps).", None),
    (-1, "King's Crossing Mall — handouts (broadcasts, journals, ham-radio transcripts).", None),
    (-1, "CMod Stack reusable component. Extract from Recruit modal; reuse in Grapple, First Impression, main Attack.", None),
    (-1, "GM force-push view to players. Switching campaign ↔ tactical or scene A ↔ B should propagate to all viewers.", None),
    (-1, "Multi-round haggling. Barter currently single-roll.", None),
    (-1, "First Impression → straight to roll modal. Skip the picker.", None),
    (-1, "GM Tools → Restore to Full Health is slow. 11 sequential UPDATEs; batch by table with .in('id', ids) + Promise.all.", None),
    (-1, "Character Evolution / CDP Calculator. Post-creation growth tool; spend earned CDP on attribute/skill/trait raises.", None),

    (1, "PRE-TESTER POLISH", None),
    (-1, "Cost-containment alarm. Supabase 75% quota + Vercel bandwidth alert. ~30 min vendor-portal config.", None),
    (-1, "Demo / sample campaign for first-time GMs. ~2-3 hours.", None),
    (-1, "Beginners' guide /welcome links. docs/beginners-guide.{txt,docx} is drafted on disk; commit + surface chapter links.", None),
    (-1, "Domain verification spot-check on Resend. FROM swap is in code; confirm outbound mail still lands.", None),
    (-1, "End-to-end smoke pass — signup → /firsttimers → /welcome → /characters/new → /map → first whisper.", None),
    (-1, "Quick Reference card on /welcome. Placeholder needs CDP / WP-RP / Stress / Inspiration cheat sheet + SRD/CRB links.", None),

    (1, "PIN / MAP", None),
    (-1, "Pin-image migration from base64 → Supabase Storage. (Character-photo migration tool exists; pin equivalent doesn't.)", None),
    (-1, "Timeline sort_order management UI. Drag-to-reorder for Thrivers; currently hardcoded via SQL.", None),

    (1, "TOOLS", None),
    (-1, "Manual crop control — drag-to-select instead of auto center-crop.", None),
    (-1, "More tools — handout generator, token template maker, roll table randomizer.", None),

    (1, "PHASE 4 (Campfire) — Tail", None),
    (-1, "Full-text search across Forums / War Stories / LFG.", None),
    (-1, "Reactions on War Stories + LFG.", None),
    (-1, "Comment threading on War Stories + LFG (Forums has it; others flat).", None),
    (-1, "Formal campaign_invitations accept/reject flow.", None),
    (-1, "LFG filters by setting + schedule.", None),
    (-1, "DZ community layer — approved player Rumors visible to all DZ campaigns.", None),

    (1, "PHASE 5 — MODULE SYSTEM", None),
    (2, "Phase D — Monetization", None),
    (-1, "Free / Paid / Premium pricing.", None),
    (-1, "Licensed GM permission unlocks paid modules.", None),
    (-1, "Author payout flow, referral tracking.", None),
    (2, "Phase E — Extras", None),
    (-1, "GM Kit Export v2 — printable PDF + module zip.", None),
    (-1, "Module + Community cross-publish.", None),
    (-1, "In-session GM toolkit — scene switcher, roster, handouts panel, roll tables linked to dice roller.", None),
    (-1, "Third-party module import (Roll20 / Foundry → Tapestry).", None),
    (2, "Phase F — GM Adventure Authoring Toolkit", None),
    (-1, "Story Arc form — guided 4-question creation surface.", None),
    (-1, "NPC quick-build inline forms.", None),
    (-1, "Map quick-build — drop new tactical scene from inside a beat.", None),
    (-1, "Handout quick-build.", None),
    (-1, "Encounter quick-build.", None),
    (-1, "Route tables — leg-by-leg encounters with roll-target each.", None),
    (-1, "Adventure preview — \"play test mode\".", None),
    (-1, "Publish Adventure — terminal step on Story Arc form.", None),

    (1, "TACTICAL MAP — Long-term", None),
    (-1, "Line of sight Phase 3 (polygon vision mask). Audit scheduled 2026-05-10.", None),
    (2, "Lv4 Skill Traits — Xero-blocked, ships together", None),
    (-1, "Inspiration Lv4 \"Beacon of Hope\" auto +4 to Morale.", None),
    (-1, "Psychology* Lv4 \"Insightful Counselor\" auto +3 to Morale.", None),
    (-1, "Generic Lv4 Trait surface on the character sheet.", None),
    (-1, "Auto-application hooks for any other Lv4 Trait.", None),
    (-1, "Barter Lv4 cheat-doubling.", None),

    (1, "CODE HEALTH", None),
    (-1, "Split table page into subcomponents. Currently 10,542 lines (was 5,365 when first deferred — it grew, didn't shrink). High risk; needs a clean day.", None),
    (-1, "Debounce realtime callbacks. Optimization-only.", None),

    (1, "DISCUSSION / UNDECIDED", None),
    (-1, "NPC health as narrative feeling. Deferred 2026-04-26; re-open if a different framing comes up.", None),
    (-1, "Decide on hide-NPCs flag. Global \"reveal to players\" boolean vs. per-instance reveal events.", None),

    (1, "TOP-LEVEL /todo.md (last updated 2026-04-11 — mostly stale)", None),
    (-1, "VERIFY + APPLY sql/initiative-order-rls-members-write.sql (Nana attack-doesn't-advance bug).", None),
    (-1, "APPLY sql/player-notes-session-tag.sql.", None),

    ("PAGEBREAK", None, None),

    (1, "LONG-TERM ROADMAP (Phases 6-11) — Aspirational, not \"open work\"", None),
    (2, "Phase 6", None),
    (-1, "LFG matching by setting + playstyle.", None),
    (-1, "Session scheduling — calendar view.", None),
    (-1, "The Gazette — auto campaign newsletter.", None),
    (-1, "Between-session experience.", None),
    (-1, "Subscriber tiers — Free / Paid / Premium.", None),
    (-1, "Graffiti — Distemper-branded reactions.", None),

    (2, "Phase 7 — Ghost Mode Advanced", None),
    (-1, "Ghost-to-Survivor funnel analytics.", None),
    (-1, "A/B test soft wall messaging.", None),
    (-1, "QR-scanner onboarding flow.", None),
    (-1, "Reactivate /firsttimers onboarding page.", None),

    (2, "Phase 8 — Physical Products", None),
    (-1, "Chased QR codes — fold-out map deep-links.", None),
    (-1, "Anonymous preview for QR scanners without accounts.", None),
    (-1, "Chased module — pre-populated with Delaware content.", None),
    (-1, "Mongrels sourcebook upload, seed pins/NPCs.", None),
    (-1, "Physical product landing pages.", None),

    (2, "Phase 9 — Maturity", None),
    (-1, "Contextual rules links from sheet + dice roller.", None),
    (-1, "Mobile optimization pass.", None),
    (-1, "Mobile dice roller.", None),
    (-1, "Global search across characters / campaigns / pins / NPCs / Campfire.", None),

    (2, "Phase 10 — Future Platforms", None),
    (-1, "Displaced — space setting on separate platform.", None),
    (-1, "Extract @xse/core monorepo.", None),
    (-1, "Each setting gets own domain + branding.", None),

    (2, "Phase 11 — Cross-Platform Parity", None),
    (-1, "Campaign Calendar — date-gated lore events.", None),
    (-1, "Roll20 Export — sheet HTML/CSS/JS, ZIP exporter, ingest.", None),

    (2, "Campaign Calendar Backburner — Revisit triggers", None),
    (-1, "Skip Week → community frozen 4+ sessions.", None),
    (-1, "World events that should've ended still applying CMod.", None),
    (-1, "\"X days passed\" → automatic ration / weather / community drift.", None),
    (-1, "Encumbrance tick auto-fire on time advancement.", None),
    (-1, "DB: campaign_clock table or jsonb on campaigns.", None),
    (-1, "Helper lib/campaign-clock.ts with advance(campaignId, hours).", None),
    (-1, "Clock widget in table page header.", None),
    (-1, "Migrate Time button from Inventory #1 to unified clock.", None),

    ("PAGEBREAK", None, None),

    (1, "SUMMARY", None),
    (0, "Top priority (playtest): 3", None),
    (0, "Older bugs (need repro / open): 5 + 5", None),
    (0, "Partials to finish: 10", None),
    (0, "UX/polish: 9", None),
    (0, "Pre-tester polish: 6", None),
    (0, "Pin/map/tools: 4", None),
    (0, "Phase 4 tail / Module Phases C-F: 22", None),
    (0, "Tactical map long-term + Lv4 (Xero-blocked): 6", None),
    (0, "Code health: 2", None),
    (0, "Discussion: 2", None),
    (0, "Stale top-level repo: 2", None),
    (0, "Long-term roadmap (Phases 6-11): ~25", None),
    (0, "", None),
    (0, "Realistic short-term: ~35-40 items not blocked on Lv4 / repro / roadmap.", None),
    (0, "For this week if shipping daily: the 3 playtest bugs + 6 pre-tester polish + 2-3 partials = a solid 11-12 item sprint.", None),
]


def build_doc(out_path: Path) -> None:
    doc = Document()

    # ── Page setup: 0.75" margins all around for tighter print ────────
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # ── Default body font: 11pt, narrow line spacing ──────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # ── Title block ───────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("The Tapestry — Open Work Checklist")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)  # tomato red

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Generated 2026-05-05 — post verification sweep")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x5A, 0x55, 0x50)
    sub.paragraph_format.space_after = Pt(12)

    # ── Iterate content ───────────────────────────────────────────────
    for entry in CONTENT:
        kind, text, italic_note = entry

        if kind == "PAGEBREAK":
            doc.add_page_break()
            continue

        if kind == 1:
            # Section header — Heading 1 styled
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            # Underline by drawing a horizontal rule via bottom border
            # (skipped — paragraph-level borders in python-docx require
            # XML hackery; the bold + size + spacing reads as a header
            # cleanly on print).
            continue

        if kind == 2:
            # Sub-section header — Heading 2 styled
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
            continue

        if kind == 0:
            # Body paragraph
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
            continue

        if kind == -1:
            # Checkbox item — "☐ <text>" with hanging indent so the
            # box sits at the left margin and the text wraps under
            # itself, not under the box.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.35)
            p.paragraph_format.space_after = Pt(4)

            box = p.add_run("☐  ")  # ☐ + two spaces
            box.font.size = Pt(13)
            box.font.name = "Cambria"  # the unicode box renders larger + clearer in serif

            body = p.add_run(text)
            body.font.size = Pt(11)

            if italic_note:
                # Italic note on a new line, indented under the item
                note_p = doc.add_paragraph()
                note_p.paragraph_format.left_indent = Inches(0.6)
                note_p.paragraph_format.space_after = Pt(6)
                note_run = note_p.add_run(italic_note)
                note_run.italic = True
                note_run.font.size = Pt(10)
                note_run.font.color.rgb = RGBColor(0x5A, 0x55, 0x50)
            continue

    # ── Footer ────────────────────────────────────────────────────────
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(20)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("end of checklist")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parent.parent
    out = repo_root / "tasks" / "open-work-2026-05-05.docx"
    build_doc(out)
